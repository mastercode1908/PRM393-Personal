import sys
import subprocess
import os

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

# Add Title
title = doc.add_heading('Onboarding UI Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Intro Text
doc.add_paragraph('This document describes the flow, purpose, and UI components of the Onboarding process for the application.')

# Add Table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Screen Name'
hdr_cells[1].text = 'Purpose'
hdr_cells[2].text = 'Key Components & Logic'

# Bold headers and set background
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    # Light gray background for header
    shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_screen_row(name, purpose, components):
    row_cells = table.add_row().cells
    
    # Name
    p_name = row_cells[0].paragraphs[0]
    run_name = p_name.add_run(name)
    run_name.bold = True
    
    # Purpose
    row_cells[1].text = purpose
    
    # Components
    row_cells[2].text = components

# Row 1
add_screen_row(
    'Splash Screen',
    'Initial screen to load resources and check authentication token.',
    '• Background: White\n'
    '• Logo: Centered icon with a purple box shadow.\n'
    '• Loading Indicator: Purple CircularProgressIndicator.\n'
    '• Logic: Waits 2 seconds (simulating local token check), then automatically redirects to Login Screen (cannot navigate back).'
)

# Row 2
add_screen_row(
    'Login Screen',
    'Provides a form for users to authenticate into the system.',
    '• AppBar: White background, "ĐĂNG NHẬP" centered title.\n'
    '• Email Input: Text field with a dynamic purple focus border.\n'
    '• Password Input: Text field with a toggleable eye icon to obscure text.\n'
    '• Login Button: Purple-to-deep-purple gradient button with shadow.\n'
    '• Logic: On press, hides the keyboard, shows a loading spinner inside the button, simulates an API call (2s delay), and redirects to the Intro Screen.'
)

# Row 3
add_screen_row(
    'Intro / Onboarding Screen',
    'Introduces key features using a swipeable image carousel.',
    '• Carousel: PageView.builder containing high-quality images, bold titles, and descriptions.\n'
    '• Bottom Navigation Row:\n'
    '   - Left: "Skip" (Bỏ qua) or "Back" (Trở lại) text button.\n'
    '   - Center: Animated dot indicators that dynamically expand for the active slide.\n'
    '   - Right: "Next" (Tiếp) or "Get Started" (Bắt đầu) solid purple button.\n'
    '• Logic: Swiping or clicking buttons changes slides. Clicking "Get Started" or "Skip" redirects the user to the Home Screen.'
)

# Row 4
add_screen_row(
    'Home Screen',
    'Final destination confirming successful onboarding and authentication.',
    '• AppBar: Purple background, "TRANG CHỦ" title.\n'
    '• Success Icon: Large purple check-circle icon inside a soft purple container.\n'
    '• Welcome Text: Confirms completion of the onboarding flow and introduces the app\'s main tools.\n'
    '• Logout Button: Red outlined button with an icon.\n'
    '• Logic: Clicking "Logout" clears the navigation stack and returns the user to the Login Screen.'
)

# Set column widths roughly
for row in table.rows:
    row.cells[0].width = Pt(100)
    row.cells[1].width = Pt(150)
    row.cells[2].width = Pt(250)

# Save document
output_path = r'd:\FUlearning\SeSummer26\PRM393\All-projects\learn_flutter\lib\OnBoarding-Screen\Onboarding_Screens_Description.docx'
doc.save(output_path)
print(f"Successfully saved Word document to: {output_path}")
